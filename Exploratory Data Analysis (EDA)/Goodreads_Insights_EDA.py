"""
=============================================================
  CodeAlpha Internship — Task 2: Exploratory Data Analysis
  Dataset : Goodreads Books Dataset (books.csv)
  Author  : CodeAlpha Intern_Yahia Nesma
=============================================================
"""

import pandas as pd
import matplotlib
matplotlib.use('Agg')          # headless rendering
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import seaborn as sns
import numpy as np
import warnings
warnings.filterwarnings('ignore')

# ─────────────────────────────────────────────────────────────
# 1. LOAD & INITIAL INSPECTION
# ─────────────────────────────────────────────────────────────
df = pd.read_csv('books.csv', on_bad_lines='skip')
df.columns = df.columns.str.strip()   # remove leading/trailing spaces
initial_rows, initial_cols = df.shape   # keep original shape for the report

print("=" * 60)
print("DATASET OVERVIEW")
print("=" * 60)
print(f"Shape      : {df.shape[0]:,} rows × {df.shape[1]} columns")
print(f"Columns    : {df.columns.tolist()}")
print(f"\nData Types:\n{df.dtypes}")
print(f"\nMissing Values:\n{df.isnull().sum()}")

# ─────────────────────────────────────────────────────────────
# 2. HANDLE MISSING DATA
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("HANDLING MISSING DATA")
print("=" * 60)

missing_before = df.isnull().sum()
missing_pct    = (missing_before / len(df) * 100).round(2)
missing_report = pd.DataFrame({'missing_count': missing_before, 'missing_pct': missing_pct})
missing_report = missing_report[missing_report['missing_count'] > 0]

if missing_report.empty:
    print("No missing values detected in any column.")
else:
    print("Columns with missing values:")
    print(missing_report.to_string())

# Strategy:
#  - categorical / text columns  -> fill with 'Unknown' (keeps the row, avoids
#    dropping books just because metadata like publisher is absent)
#  - core numeric columns required for analysis -> drop the row, since an
#    imputed rating/page-count/author would distort every downstream stat
for cat_col in ['publisher', 'language_code', 'authors', 'title']:
    if cat_col in df.columns and df[cat_col].isnull().any():
        n_filled = df[cat_col].isnull().sum()
        df[cat_col] = df[cat_col].fillna('Unknown')
        print(f"  → Filled {n_filled} missing '{cat_col}' values with 'Unknown'")

essential_cols = [c for c in ['average_rating', 'num_pages', 'ratings_count',
                               'text_reviews_count', 'publication_date'] if c in df.columns]
rows_before = len(df)
df = df.dropna(subset=essential_cols)
rows_dropped = rows_before - len(df)
print(f"  → Dropped {rows_dropped} rows missing essential numeric fields {essential_cols}")

print(f"\nMissing values remaining:\n{df.isnull().sum().sum()} total cells")

# ─────────────────────────────────────────────────────────────
# 3. CLEANING (outlier / invalid-value filtering)
# ─────────────────────────────────────────────────────────────
df = df[df['num_pages'] > 0]          # drop books with 0 pages
df = df[df['average_rating'] > 0]     # drop books with 0 rating

# Parse publication year
df['pub_year'] = pd.to_datetime(df['publication_date'], errors='coerce').dt.year

# Log-transform ratings_count for scatter plots
df['ratings_count_log'] = np.log1p(df['ratings_count'])

print(f"\nAfter cleaning : {df.shape[0]:,} rows")

# ─────────────────────────────────────────────────────────────
# 3. UNIVARIATE ANALYSIS
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("UNIVARIATE ANALYSIS")
print("=" * 60)

print("\n--- Average Rating ---")
print(df['average_rating'].describe().round(3))

print("\n--- Page Count ---")
print(df['num_pages'].describe().round(1))

print("\n--- Ratings Count ---")
print(df['ratings_count'].describe().round(0))

print("\n--- Top 10 Authors by Book Count ---")
print(df['authors'].value_counts().head(10).to_string())

print("\n--- Top 10 Publishers ---")
print(df['publisher'].value_counts().head(10).to_string())

print("\n--- Language Distribution ---")
print(df['language_code'].value_counts().to_string())

# ─────────────────────────────────────────────────────────────
# 4. BIVARIATE / CORRELATION ANALYSIS
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("CORRELATION ANALYSIS")
print("=" * 60)
corr_cols = ['average_rating', 'num_pages', 'ratings_count', 'text_reviews_count']
print(df[corr_cols].corr().round(3).to_string())

# ─────────────────────────────────────────────────────────────
# 5. VISUALISATIONS  (11 charts on one dashboard)
# ─────────────────────────────────────────────────────────────
palette   = ['#2D6A4F','#52B788','#95D5B2','#D8F3DC','#B7E4C7',
             '#1B4332','#40916C','#74C69D','#081C15','#74C69D']
ax_bg     = '#152E21'
text_col  = '#B7E4C7'
grid_col  = '#2D4A38'

sns.set_style("whitegrid")
plt.rcParams.update({'font.family': 'DejaVu Sans',
                     'axes.spines.top': False,
                     'axes.spines.right': False})

def style_ax(ax):
    ax.set_facecolor(ax_bg)
    ax.tick_params(colors=text_col)
    ax.xaxis.label.set_color(text_col)
    ax.yaxis.label.set_color(text_col)
    for sp in ax.spines.values():
        sp.set_color(grid_col)
    ax.grid(color=grid_col, linewidth=0.5, alpha=0.7)

title_kw = dict(fontsize=11, fontweight='bold', color='#D8F3DC', pad=10)

fig = plt.figure(figsize=(22, 28))
fig.patch.set_facecolor('#0F1B14')
gs  = gridspec.GridSpec(4, 3, figure=fig, hspace=0.55, wspace=0.38)

# ── Chart 1 : Rating Distribution ────────────────────────────
ax1 = fig.add_subplot(gs[0, 0]); style_ax(ax1)
n, bins, patches = ax1.hist(df['average_rating'], bins=40, edgecolor='#081C15', lw=0.4)
for patch, val in zip(patches, bins):
    patch.set_facecolor(plt.cm.YlGn(0.2 + 0.7 * (val - 2) / 3))
ax1.axvline(df['average_rating'].mean(),   color='#FF6B6B', ls='--', lw=1.5,
            label=f'Mean: {df["average_rating"].mean():.2f}')
ax1.axvline(df['average_rating'].median(), color='#FFD93D', ls='--', lw=1.5,
            label=f'Median: {df["average_rating"].median():.2f}')
ax1.legend(fontsize=8, facecolor=ax_bg, labelcolor=text_col, edgecolor=grid_col)
ax1.set_title('Rating Distribution', **title_kw)
ax1.set_xlabel('Average Rating'); ax1.set_ylabel('Number of Books')

# ── Chart 2 : Top 10 Authors ──────────────────────────────────
ax2 = fig.add_subplot(gs[0, 1]); style_ax(ax2)
top_authors = df['authors'].value_counts().head(10)
c2 = sns.color_palette('YlGn', 10)[::-1]
bars2 = ax2.barh(top_authors.index[::-1], top_authors.values[::-1], color=c2,
                 edgecolor='#081C15', lw=0.4)
for bar, val in zip(bars2, top_authors.values[::-1]):
    ax2.text(bar.get_width() + .3, bar.get_y() + bar.get_height()/2,
             str(val), va='center', fontsize=8, color=text_col)
ax2.set_title('Top 10 Authors by Book Count', **title_kw)
ax2.set_xlabel('Number of Books')

# ── Chart 3 : Top 10 Publishers ──────────────────────────────
ax3 = fig.add_subplot(gs[0, 2]); style_ax(ax3)
top_pub = df['publisher'].value_counts().head(10)
ax3.bar(range(len(top_pub)), top_pub.values, color=sns.color_palette('YlGn', 10)[::-1],
        edgecolor='#081C15', lw=0.4)
ax3.set_xticks(range(len(top_pub)))
ax3.set_xticklabels(top_pub.index, rotation=40, ha='right', fontsize=7.5)
ax3.set_title('Top 10 Publishers', **title_kw)
ax3.set_ylabel('Number of Books')

# ── Chart 4 : Language Pie ─────────────────────────────────────
ax4 = fig.add_subplot(gs[1, 0]); ax4.set_facecolor(ax_bg)
lang = df['language_code'].value_counts()
pie_d = list(lang.head(6).values) + [lang.iloc[6:].sum()]
pie_l = list(lang.head(6).index)  + ['Other']
ax4.pie(pie_d, labels=pie_l, autopct='%1.1f%%',
        colors=sns.color_palette('YlGn', len(pie_d)),
        explode=[0.04]*len(pie_d), startangle=140,
        textprops={'color': text_col, 'fontsize': 8},
        wedgeprops={'edgecolor': ax_bg, 'linewidth': 1.2})
ax4.set_title('Language Distribution', **title_kw)

# ── Chart 5 : Page Count Violin ───────────────────────────────
ax5 = fig.add_subplot(gs[1, 1]); style_ax(ax5)
pages_c = df['num_pages'].clip(upper=df['num_pages'].quantile(0.99))
vp = ax5.violinplot(pages_c, positions=[0], showmedians=True, showextrema=True)
for pc in vp['bodies']:
    pc.set_facecolor('#52B788'); pc.set_edgecolor('#2D6A4F'); pc.set_alpha(0.75)
vp['cmedians'].set_color('#FFD93D')
for k in ('cmins','cmaxes','cbars'):
    vp[k].set_color(text_col)
ax5.set_xticks([0]); ax5.set_xticklabels(['Books'])
ax5.set_title('Page Count Distribution\n(clipped at 99th pct)', **title_kw)
ax5.set_ylabel('Number of Pages')
for y, lbl, c in [(pages_c.median(), f'Median: {pages_c.median():.0f}', '#FFD93D'),
                  (pages_c.mean(),   f'Mean: {pages_c.mean():.0f}',   '#FF6B6B')]:
    ax5.axhline(y, color=c, ls='--', lw=1)
    ax5.text(0.52, y, lbl, va='center', fontsize=8, color=c)

# ── Chart 6 : Scatter Ratings Count vs Rating ─────────────────
ax6 = fig.add_subplot(gs[1, 2]); style_ax(ax6)
sc = ax6.scatter(df['ratings_count_log'], df['average_rating'],
                 alpha=0.25, s=8, c=df['average_rating'], cmap='YlGn', linewidths=0)
cb = plt.colorbar(sc, ax=ax6)
cb.set_label('Avg Rating', color=text_col, fontsize=8)
cb.ax.yaxis.set_tick_params(color=text_col)
plt.setp(cb.ax.yaxis.get_ticklabels(), color=text_col, fontsize=7)
cb.outline.set_edgecolor(grid_col)
ax6.set_title('Ratings Count vs Average Rating', **title_kw)
ax6.set_xlabel('Log(Ratings Count)'); ax6.set_ylabel('Average Rating')

# ── Chart 7 : Publications Per Year ──────────────────────────
ax7 = fig.add_subplot(gs[2, :2]); style_ax(ax7)
yearly = df[df['pub_year'].between(1980, 2010)]['pub_year'].value_counts().sort_index()
ax7.fill_between(yearly.index, yearly.values, alpha=0.35, color='#52B788')
ax7.plot(yearly.index, yearly.values, color='#52B788', lw=2)
ax7.scatter(yearly.index, yearly.values, color='#95D5B2', s=25, zorder=5)
peak = yearly.idxmax()
ax7.annotate(f'Peak: {int(peak)}', xy=(peak, yearly.max()),
             xytext=(peak-5, yearly.max()*0.85),
             arrowprops=dict(arrowstyle='->', color='#FFD93D'),
             color='#FFD93D', fontsize=9)
ax7.set_title('Books Published Per Year (1980–2010)', **title_kw)
ax7.set_xlabel('Year'); ax7.set_ylabel('Number of Books')

# ── Chart 8 : Avg Rating by Language ─────────────────────────
ax8 = fig.add_subplot(gs[2, 2]); style_ax(ax8)
top6 = df['language_code'].value_counts().head(6).index
lr   = df[df['language_code'].isin(top6)].groupby('language_code')['average_rating'].mean().sort_values()
bars8 = ax8.barh(lr.index, lr.values, color=sns.color_palette('YlGn', len(lr)),
                 edgecolor='#081C15', lw=0.4)
ax8.set_xlim(3.5, 4.3)
for bar, val in zip(bars8, lr.values):
    ax8.text(val+.005, bar.get_y()+bar.get_height()/2,
             f'{val:.2f}', va='center', fontsize=8, color=text_col)
ax8.set_title('Avg Rating by Language (Top 6)', **title_kw)
ax8.set_xlabel('Average Rating')

# ── Chart 9 : Correlation Heatmap ─────────────────────────────
ax9 = fig.add_subplot(gs[3, 0]); ax9.set_facecolor(ax_bg)
lbls = ['Avg Rating','Pages','Ratings','Reviews']
corr = df[corr_cols].corr()
im   = ax9.imshow(corr, cmap='YlGn', aspect='auto', vmin=-0.1, vmax=1)
plt.colorbar(im, ax=ax9, shrink=0.8).ax.tick_params(colors=text_col, labelsize=7)
ax9.set_xticks(range(4)); ax9.set_yticks(range(4))
ax9.set_xticklabels(lbls, rotation=30, ha='right', color=text_col, fontsize=8)
ax9.set_yticklabels(lbls, color=text_col, fontsize=8)
for i in range(4):
    for j in range(4):
        ax9.text(j, i, f'{corr.iloc[i,j]:.2f}', ha='center', va='center', fontsize=8,
                 color='#081C15' if corr.iloc[i,j] > 0.5 else text_col)
ax9.set_title('Correlation Heatmap', **title_kw)

# ── Chart 10 : Avg Rating by Page Bin ─────────────────────────
ax10 = fig.add_subplot(gs[3, 1]); style_ax(ax10)
bins_pg  = [0,100,200,300,400,500,700,1000,7000]
labels_pg= ['<100','100-200','200-300','300-400','400-500','500-700','700-1k','>1k']
df['page_bin'] = pd.cut(df['num_pages'], bins=bins_pg, labels=labels_pg)
pg_rat = df.groupby('page_bin', observed=True)['average_rating'].mean()
ax10.bar(range(len(pg_rat)), pg_rat.values, color=sns.color_palette('YlGn', len(pg_rat)),
         edgecolor='#081C15', lw=0.4)
ax10.set_xticks(range(len(pg_rat)))
ax10.set_xticklabels(labels_pg, rotation=35, ha='right', fontsize=8)
ax10.set_ylim(3.5, 4.3)
ax10.set_title('Avg Rating by Page Count Range', **title_kw)
ax10.set_ylabel('Average Rating')

# ── Chart 11 : Top 10 Most-Rated Books ───────────────────────
ax11 = fig.add_subplot(gs[3, 2]); style_ax(ax11)
top_r = df.nlargest(10, 'ratings_count')[['title','ratings_count']].reset_index(drop=True)
top_r['short'] = top_r['title'].str[:30] + '…'
ax11.barh(top_r['short'][::-1], (top_r['ratings_count']/1e6)[::-1],
          color=sns.color_palette('YlGn', 10), edgecolor='#081C15', lw=0.4)
ax11.set_title('Top 10 Most-Rated Books (M)', **title_kw)
ax11.set_xlabel('Ratings (millions)')

# ── Super-title & save ────────────────────────────────────────
fig.suptitle('📚  Goodreads Books Dataset — Exploratory Data Analysis',
             fontsize=18, fontweight='bold', color='#D8F3DC', y=0.98)

output_path = 'goodreads_eda.png'
plt.savefig(output_path, dpi=150, bbox_inches='tight',
            facecolor=fig.get_facecolor())
print(f"✅  Dashboard saved to '{output_path}'")

# ─────────────────────────────────────────────────────────────
# 6. KEY FINDINGS (printed summary)
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("KEY FINDINGS")
print("=" * 60)
print(f"1. Total books analysed       : {len(df):,}")
print(f"2. Average rating             : {df['average_rating'].mean():.2f}  (std={df['average_rating'].std():.2f})")
print(f"3. Median page count          : {df['num_pages'].median():.0f} pages")
print(f"4. Most prolific author       : {df['authors'].value_counts().idxmax()} ({df['authors'].value_counts().max()} books)")
print(f"5. Top publisher              : {df['publisher'].value_counts().idxmax()} ({df['publisher'].value_counts().max()} books)")
print(f"6. Dominant language          : eng/en-US = {df['language_code'].isin(['eng','en-US']).mean()*100:.1f}% of books")
print(f"7. Ratings vs Reviews corr.   : {df['ratings_count'].corr(df['text_reviews_count']):.3f} (strong positive)")
print(f"8. Rating vs Pages corr.      : {df['average_rating'].corr(df['num_pages']):.3f} (weak positive)")
print(f"9. Most-rated book            : {df.nlargest(1,'ratings_count')['title'].values[0][:60]}")
print(f"   with {df['ratings_count'].max():,} ratings")

# ─────────────────────────────────────────────────────────────
# 7. PLAIN-ENGLISH OBSERVATIONS
# ─────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("OBSERVATIONS (PLAIN ENGLISH)")
print("=" * 60)
dom_lang_pct = df['language_code'].isin(['eng', 'en-US', 'en-GB', 'en-CA']).mean() * 100
rating_review_corr = df['ratings_count'].corr(df['text_reviews_count'])
rating_pages_corr  = df['average_rating'].corr(df['num_pages'])
print(f"- Most books in this dataset are rated between "
      f"{df['average_rating'].quantile(0.25):.1f} and {df['average_rating'].quantile(0.75):.1f} stars, "
      f"clustering tightly around a mean of {df['average_rating'].mean():.2f}.")
print(f"- English-language editions dominate the catalogue, making up {dom_lang_pct:.1f}% of all books.")
print(f"- Books with more ratings tend to also have more written reviews "
      f"(correlation = {rating_review_corr:.2f}), suggesting engaged readers do both.")
print(f"- Page count has only a weak relationship with average rating "
      f"(correlation = {rating_pages_corr:.2f}) — longer books are not reliably rated higher or lower.")
print(f"- {df['authors'].value_counts().idxmax()} is the most prolific author in the dataset, "
      f"with {df['authors'].value_counts().max()} titles listed.")

# ─────────────────────────────────────────────────────────────
# 8. EXPORT POLISHED WORD REPORT (.docx)
# ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN_RGB      = RGBColor(0x2D, 0x6A, 0x4F)
DARKGREEN_RGB  = RGBColor(0x1B, 0x43, 0x32)
GREY_RGB       = RGBColor(0x59, 0x59, 0x59)
HEADER_FILL    = "2D6A4F"
ROW_FILL       = "EAF6EE"

def shade_cell(cell, hex_color):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, align_center=False):
    cell.text = ""
    para = cell.paragraphs[0]
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = DARKGREEN_RGB if level == 1 else GREEN_RGB
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align_center=False):
    para = doc.add_paragraph()
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = color
    return para

def add_bullet(doc, text, bold=False):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    return para

report_doc = Document()

# Base font
report_doc.styles['Normal'].font.name = 'Arial'
report_doc.styles['Normal'].font.size = Pt(11)

# ── Title block ───────────────────────────────────────────────
add_para(report_doc, "Exploratory Data Analysis Report", bold=True, size=26,
          color=DARKGREEN_RGB, align_center=True)
add_para(report_doc, "Goodreads Books Dataset", italic=True, size=14,
          color=GREEN_RGB, align_center=True)
add_para(report_doc, "CodeAlpha Internship — Task 2  |  Author: Nesma Yahia", size=10,
          color=GREY_RGB, align_center=True)

# ── 1. Dataset Description ──────────────────────────────────
add_heading(report_doc, "1. Dataset Description", level=1)
add_para(report_doc,
    f"The dataset used for this analysis is the Goodreads Books dataset (books.csv), "
    f"a collection of book metadata scraped from Goodreads.com. It contains "
    f"{initial_rows:,} records and {initial_cols} columns before cleaning, covering "
    f"identifying information, authorship, ratings, and publication details.")
add_heading(report_doc, "Columns", level=2)
for line in [
    "bookID — unique identifier for each book",
    "title, authors — book title and author(s), multiple authors separated by '/'",
    "average_rating — mean Goodreads user rating (0–5 scale)",
    "isbn, isbn13 — book identifiers",
    "language_code — primary language of the edition",
    "num_pages — page count",
    "ratings_count, text_reviews_count — number of ratings and written reviews received",
    "publication_date, publisher — release date and publishing house",
]:
    add_bullet(report_doc, line)

# ── 2. Data Cleaning & Missing Data Handling ─────────────────
add_heading(report_doc, "2. Data Cleaning & Missing Data Handling", level=1)
if missing_report.empty:
    missing_summary = "no missing values were present in this run of the dataset"
else:
    missing_summary = (f"missing values were found in {len(missing_report)} column(s): "
                        f"{', '.join(missing_report.index.tolist())}")
add_para(report_doc,
    f"Before analysis, the dataset was inspected for missing values across all "
    f"{initial_cols} columns using df.isnull().sum(); {missing_summary}. The pipeline "
    f"includes an explicit handling step so it remains robust against incomplete data "
    f"in other versions or future updates of the dataset:")
add_bullet(report_doc,
    "Categorical/text columns (publisher, language_code, authors, title): missing "
    "values are filled with 'Unknown' rather than dropped, preserving the row since a "
    "missing publisher or language doesn't affect the rest of a book's metrics.")
add_bullet(report_doc,
    f"Essential numeric columns ({', '.join(essential_cols)}): rows missing any of "
    f"these are dropped ({rows_dropped} rows removed this run), since imputing a "
    f"rating or page count would distort downstream statistics and correlations.")
add_para(report_doc,
    f"After missing-data handling, additional cleaning removed invalid records — books "
    f"listed with 0 pages or a 0.0 average rating (placeholder values rather than real "
    f"measurements) — leaving {len(df):,} books for analysis.")

# ── 3. Key Statistics ─────────────────────────────────────────
add_heading(report_doc, "3. Key Statistics", level=1)
stats_rows = [
    ("Total books analysed", f"{len(df):,} (after cleaning)"),
    ("Average rating", f"{df['average_rating'].mean():.2f} / 5  (std. dev. = {df['average_rating'].std():.2f})"),
    ("Median page count", f"{df['num_pages'].median():.0f} pages"),
    ("Most prolific author", f"{df['authors'].value_counts().idxmax()} — {df['authors'].value_counts().max()} books"),
    ("Top publisher", f"{df['publisher'].value_counts().idxmax()} — {df['publisher'].value_counts().max()} books"),
    ("Dominant language", f"English (eng / en-US) — {df['language_code'].isin(['eng','en-US']).mean()*100:.1f}% of books"),
    ("Ratings vs. text reviews correlation", f"{rating_review_corr:.3f}  ({'strong' if rating_review_corr > 0.5 else 'weak'} positive)"),
    ("Rating vs. page count correlation", f"{rating_pages_corr:.3f}  ({'strong' if rating_pages_corr > 0.5 else 'weak'} positive)"),
    ("Most-rated book", f"{df.nlargest(1,'ratings_count')['title'].values[0][:60]} — {df['ratings_count'].max():,} ratings"),
]
table = report_doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
for i, (label, value) in enumerate(stats_rows):
    row = table.add_row()
    set_cell_text(row.cells[0], label, bold=True)
    set_cell_text(row.cells[1], value)
    if i % 2 == 0:
        shade_cell(row.cells[0], ROW_FILL)
        shade_cell(row.cells[1], ROW_FILL)

add_para(report_doc, "")
add_heading(report_doc, "Correlation Matrix", level=2)
add_para(report_doc, "Pearson correlation coefficients between the four core numeric variables:")
corr_labels = ['Avg Rating', 'Pages', 'Ratings', 'Reviews']
corr_table = report_doc.add_table(rows=0, cols=5)
corr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
corr_table.style = 'Table Grid'
header_row = corr_table.add_row()
set_cell_text(header_row.cells[0], "", bold=True)
for j, lbl in enumerate(corr_labels):
    set_cell_text(header_row.cells[j + 1], lbl, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), align_center=True)
    shade_cell(header_row.cells[j + 1], HEADER_FILL)
shade_cell(header_row.cells[0], HEADER_FILL)
for i, row_label in enumerate(corr_labels):
    row = corr_table.add_row()
    set_cell_text(row.cells[0], row_label, bold=True)
    shade_cell(row.cells[0], ROW_FILL)
    for j in range(4):
        set_cell_text(row.cells[j + 1], f"{corr.iloc[i, j]:.3f}", align_center=True)

# ── 4. Visual Dashboard ──────────────────────────────────────
report_doc.add_page_break()
add_heading(report_doc, "4. Visual Dashboard", level=1)
add_para(report_doc,
    "The figure below summarises eleven charts covering rating distribution, top "
    "authors/publishers, language mix, page-count distribution, ratings-vs-rating "
    "relationships, publication trends, and correlation structure.")
img_para = report_doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_para.add_run().add_picture(output_path, width=Inches(5.7))

# ── 5. Findings & Observations ───────────────────────────────
report_doc.add_page_break()
add_heading(report_doc, "5. Findings & Observations", level=1)
findings = [
    f"Most books in this dataset are rated between {df['average_rating'].quantile(0.25):.1f} and "
    f"{df['average_rating'].quantile(0.75):.1f} stars, clustering tightly around a mean of "
    f"{df['average_rating'].mean():.2f} — Goodreads ratings skew positive overall.",
    f"English-language editions dominate the catalogue, making up {dom_lang_pct:.1f}% of all "
    f"books (eng + en-US/en-GB/en-CA), with other languages trailing far behind.",
    f"Ratings count and text-review count move together strongly (correlation = "
    f"{rating_review_corr:.2f}): books that attract many star ratings also attract many "
    f"written reviews, suggesting the same engaged readers drive both.",
    f"Page count has only a weak relationship with average rating (correlation = "
    f"{rating_pages_corr:.2f}) — longer books are not reliably rated higher or lower than "
    f"shorter ones.",
    f"{df['authors'].value_counts().idxmax()} is the most prolific author in the dataset, "
    f"with {df['authors'].value_counts().max()} titles listed.",
    f"Book publication volume rose sharply over the years and peaked around "
    f"{int(peak)}, before dropping off — most likely reflecting when this dataset "
    f"snapshot was scraped rather than an actual decline in publishing.",
    f"{df.nlargest(1,'ratings_count')['title'].values[0][:60]} is the single most-rated "
    f"book in the dataset by a wide margin, with {df['ratings_count'].max():,} ratings.",
]
for f_text in findings:
    add_bullet(report_doc, f_text, bold=True)

# ── 6. Conclusion ─────────────────────────────────────────────
add_heading(report_doc, "6. Conclusion", level=1)
add_para(report_doc,
    f"This exploratory analysis shows a dataset dominated by English-language fiction "
    f"with generally positive reception (mean rating ≈ {df['average_rating'].mean():.2f}/5). "
    f"Reader engagement metrics (ratings and reviews) are tightly linked, while book length "
    f"is only weakly tied to perceived quality. A handful of prolific authors and major "
    f"publishers account for a disproportionate share of titles, and a small number of "
    f"breakout books capture the vast majority of reader attention. These patterns are "
    f"typical of popularity-driven platforms, where a long tail of moderately-read books "
    f"coexists with a few outsized bestsellers.")

report_path = 'Goodreads_EDA_Report.docx'
report_doc.save(report_path)

# python-docx's default template emits <w:zoom w:val="bestFit"/> without the
# required w:percent attribute, which some strict validators reject. Patch it
# in-place so the file is fully schema-valid.
import zipfile, shutil, tempfile

def fix_zoom_attribute(docx_path):
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = data.replace(b'<w:zoom w:val="bestFit"/>',
                                     b'<w:zoom w:val="bestFit" w:percent="100"/>')
            zout.writestr(item, data)
    shutil.move(tmp_path, docx_path)

fix_zoom_attribute(report_path)
print(f"✅  Word report saved to '{report_path}'")
