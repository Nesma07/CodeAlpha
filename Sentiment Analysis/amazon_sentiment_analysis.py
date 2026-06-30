"""
=============================================================
  Sentiment Analysis — Amazon Product Reviews
  Goal: Automatically label reviews as positive, negative, or
        neutral, detect specific emotions, and surface the
        patterns that matter for product/marketing decisions.
=============================================================

This script:
  1. Loads a real-world text dataset (Amazon product reviews —
     review text + star rating + product info)
  2. Cleans / preprocesses the review text
  3. Runs lexicon-based sentiment analysis (VADER) to classify
     each review as positive, negative, or neutral
  4. Runs a simple emotion lexicon (NRC-style word lists) to
     detect specific emotions (joy, anger, trust, etc.)
  5. Visualizes the results (sentiment counts, emotion mix,
     sentiment vs. star rating, sentiment by product/brand)
  6. Builds a polished Word report explaining the method and
     the findings in plain English
"""

import re
import json
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import warnings
warnings.filterwarnings('ignore')

sns.set_style("whitegrid")
plt.rcParams.update({'font.family': 'DejaVu Sans',
                      'axes.spines.top': False, 'axes.spines.right': False})
COLORS = {'Positive': '#2D6A4F', 'Neutral': '#FFB100', 'Negative': '#D62828'}
TITLE_KW = dict(fontsize=14, fontweight='bold', color='#1B4332', pad=14)

# ─────────────────────────────────────────────────────────────
# 1. LOAD THE TEXT DATASET
# ─────────────────────────────────────────────────────────────
# Real Amazon product reviews (electronics/tablets category):
# review text, star rating, product name, brand, category.
df = pd.read_csv('amazon_reviews.csv',
                  usecols=['name', 'brand', 'categories',
                           'reviews.rating', 'reviews.text', 'reviews.title'])
df = df.rename(columns={'reviews.rating': 'rating',
                         'reviews.text': 'review_text',
                         'reviews.title': 'review_title'})

print("=" * 60)
print("DATASET OVERVIEW")
print("=" * 60)
print(f"Raw reviews loaded : {len(df):,}")
print(f"Columns            : {df.columns.tolist()}")
print(f"Missing review text: {df['review_text'].isnull().sum()}")
print(f"Missing rating     : {df['rating'].isnull().sum()}")

# ─────────────────────────────────────────────────────────────
# 2. CLEAN / PREPROCESS THE TEXT
# ─────────────────────────────────────────────────────────────
# Drop rows with no review text or no rating — can't analyze or
# validate sentiment without them.
df = df.dropna(subset=['review_text', 'rating'])
df['name'] = df['name'].fillna('Unknown product')

def clean_product_name(name):
    """Some product names are duplicated with a newline (e.g. 'Echo (White),,,\\nEcho (White),,,').
    Collapse those, strip stray commas/whitespace."""
    name = str(name)
    if '\n' in name:
        parts = [p.strip() for p in name.split('\n') if p.strip()]
        if len(parts) > 1 and parts[0] == parts[1]:
            name = parts[0]
        else:
            name = parts[0] if parts else name
    name = re.sub(r'(,\s*){2,}', ', ', name).strip(', ').strip()
    return name

df['name'] = df['name'].apply(clean_product_name)

def clean_text(text):
    """Light cleaning: keep sentiment-bearing punctuation (VADER
    actually uses '!' and CAPS as signals), just strip whitespace,
    URLs, and HTML leftovers."""
    text = str(text)
    text = re.sub(r'http\S+|www\.\S+', '', text)      # URLs
    text = re.sub(r'<.*?>', '', text)                  # HTML tags
    text = re.sub(r'\s+', ' ', text).strip()            # extra whitespace
    return text

df['clean_text'] = df['review_text'].apply(clean_text)
df = df[df['clean_text'].str.len() > 0]
df['word_count'] = df['clean_text'].str.split().str.len()

print(f"\nAfter cleaning     : {len(df):,} reviews ready for analysis")
print(f"Average word count : {df['word_count'].mean():.1f} words/review")

# ─────────────────────────────────────────────────────────────
# 3. SENTIMENT ANALYSIS (lexicon-based — VADER)
# ─────────────────────────────────────────────────────────────
# VADER is tuned for short, informal text like reviews/social
# media — it understands punctuation, capitalization, and
# negation ("not good") without needing model training.
analyzer = SentimentIntensityAnalyzer()

def get_sentiment(text):
    scores = analyzer.polarity_scores(text)
    compound = scores['compound']
    if compound >= 0.05:
        label = 'Positive'
    elif compound <= -0.05:
        label = 'Negative'
    else:
        label = 'Neutral'
    return pd.Series([compound, label])

df[['sentiment_score', 'sentiment_label']] = df['clean_text'].apply(get_sentiment)

print("\n" + "=" * 60)
print("SENTIMENT CLASSIFICATION")
print("=" * 60)
sent_counts = df['sentiment_label'].value_counts()
sent_pct = (sent_counts / len(df) * 100).round(1)
for label in ['Positive', 'Neutral', 'Negative']:
    if label in sent_counts:
        print(f"{label:9s}: {sent_counts[label]:>6,}  ({sent_pct[label]}%)")

# Cross-check: does sentiment label agree with star rating?
# (validates the lexicon approach against ground-truth ratings)
df['rating_sentiment'] = pd.cut(df['rating'], bins=[0, 2, 3, 5],
                                  labels=['Negative', 'Neutral', 'Positive'])
agreement = (df['sentiment_label'] == df['rating_sentiment']).mean() * 100
print(f"\nAgreement between text sentiment and star rating: {agreement:.1f}%")

# ─────────────────────────────────────────────────────────────
# 4. EMOTION DETECTION (simple lexicon-based approach)
# ─────────────────────────────────────────────────────────────
# A compact NRC-style emotion lexicon: maps words to specific
# emotions, not just positive/negative polarity.
EMOTION_LEXICON = {
    'joy':      ['love', 'happy', 'great', 'excellent', 'amazing', 'awesome',
                 'wonderful', 'fantastic', 'perfect', 'delight', 'pleased',
                 'enjoy', 'glad', 'satisfied', 'best'],
    'trust':    ['reliable', 'trust', 'recommend', 'quality', 'durable',
                 'consistent', 'dependable', 'confident', 'secure', 'sturdy'],
    'anger':    ['angry', 'furious', 'annoyed', 'frustrat', 'hate', 'terrible',
                 'awful', 'worst', 'rage', 'mad', 'irritat'],
    'sadness':  ['disappoint', 'sad', 'unhappy', 'regret', 'unfortunate',
                 'poor', 'sorry', 'letdown', 'fail'],
    'fear':     ['worried', 'concern', 'afraid', 'risk', 'unsafe', 'danger',
                 'scared', 'anxious', 'doubt'],
    'surprise': ['surprised', 'shocked', 'unexpected', 'wow', 'surprising',
                 'astonished', 'amazed'],
}

def detect_emotions(text):
    text_lower = text.lower()
    found = []
    for emotion, words in EMOTION_LEXICON.items():
        if any(w in text_lower for w in words):
            found.append(emotion)
    return found if found else ['none']

df['emotions'] = df['clean_text'].apply(detect_emotions)
all_emotions = [e for sub in df['emotions'] for e in sub if e != 'none']
emotion_counts = Counter(all_emotions)

print("\n" + "=" * 60)
print("EMOTION DETECTION")
print("=" * 60)
for emo, cnt in emotion_counts.most_common():
    print(f"{emo.capitalize():10s}: {cnt:>6,} mentions")

# ─────────────────────────────────────────────────────────────
# 5. CLASSIFY & AGGREGATE BY PRODUCT
# ─────────────────────────────────────────────────────────────
product_sentiment = (df.groupby('name')['sentiment_label']
                       .value_counts(normalize=True)
                       .unstack(fill_value=0) * 100).round(1)
top_products = df['name'].value_counts().head(8).index
product_sentiment_top = product_sentiment.loc[top_products]

print("\n" + "=" * 60)
print("SENTIMENT BY PRODUCT (top 8 most-reviewed)")
print("=" * 60)
print(product_sentiment_top.to_string())

# ─────────────────────────────────────────────────────────────
# 6. VISUALIZE THE RESULTS
# ─────────────────────────────────────────────────────────────

# Chart 1: Overall sentiment counts (bar chart)
fig, ax = plt.subplots(figsize=(8, 5.5))
order = ['Positive', 'Neutral', 'Negative']
counts_ordered = [sent_counts.get(l, 0) for l in order]
bars = ax.bar(order, counts_ordered, color=[COLORS[l] for l in order],
              edgecolor='white', linewidth=1.2, width=0.6)
for bar, val, pct in zip(bars, counts_ordered, [sent_pct.get(l, 0) for l in order]):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(counts_ordered)*0.01,
            f'{val:,}\n({pct}%)', ha='center', fontsize=11, fontweight='bold')
ax.set_title('Overall Sentiment of Amazon Reviews', **TITLE_KW)
ax.set_ylabel('Number of Reviews', fontsize=11)
plt.tight_layout()
plt.savefig('chart_1_sentiment_counts.png', dpi=150, bbox_inches='tight')
plt.close()

# Chart 2: Sentiment distribution as donut/pie
fig, ax = plt.subplots(figsize=(7, 7))
ax.pie(counts_ordered, labels=order, autopct='%1.1f%%', startangle=90,
       colors=[COLORS[l] for l in order],
       wedgeprops={'edgecolor': 'white', 'linewidth': 2, 'width': 0.45},
       textprops={'fontsize': 12, 'fontweight': 'bold'})
ax.set_title('Sentiment Share of All Reviews', **TITLE_KW)
plt.tight_layout()
plt.savefig('chart_2_sentiment_share.png', dpi=150, bbox_inches='tight')
plt.close()

# Chart 3: Emotion frequency (bar chart)
fig, ax = plt.subplots(figsize=(9, 5.5))
emo_sorted = emotion_counts.most_common()
emo_labels = [e[0].capitalize() for e in emo_sorted]
emo_values = [e[1] for e in emo_sorted]
bars = ax.barh(emo_labels[::-1], emo_values[::-1],
               color=sns.color_palette('YlGn', len(emo_labels))[::-1],
               edgecolor='#1B4332', linewidth=0.6)
for bar, val in zip(bars, emo_values[::-1]):
    ax.text(bar.get_width() + max(emo_values)*0.01, bar.get_y() + bar.get_height()/2,
            f'{val:,}', va='center', fontsize=10)
ax.set_title('Specific Emotions Detected in Reviews', **TITLE_KW)
ax.set_xlabel('Number of Mentions', fontsize=11)
plt.tight_layout()
plt.savefig('chart_3_emotion_counts.png', dpi=150, bbox_inches='tight')
plt.close()

# Chart 4: Sentiment vs. star rating (validates the method)
fig, ax = plt.subplots(figsize=(9, 5.5))
cross = pd.crosstab(df['rating'], df['sentiment_label'])
cross = cross.reindex(columns=order, fill_value=0)
cross.plot(kind='bar', stacked=True, ax=ax,
           color=[COLORS[l] for l in order], edgecolor='white', linewidth=0.5)
ax.set_title('Text Sentiment vs. Star Rating', **TITLE_KW)
ax.set_xlabel('Star Rating', fontsize=11); ax.set_ylabel('Number of Reviews', fontsize=11)
ax.legend(title='Sentiment', fontsize=10)
plt.xticks(rotation=0)
plt.tight_layout()
plt.savefig('chart_4_sentiment_vs_rating.png', dpi=150, bbox_inches='tight')
plt.close()

# Chart 5: Sentiment mix by top products (heatmap)
fig, ax = plt.subplots(figsize=(9, 6))
sns.heatmap(product_sentiment_top[order], annot=True, fmt='.0f', cmap='YlGn',
            linewidths=1, linecolor='white', cbar_kws={'label': '% of reviews'}, ax=ax)
ax.set_title('Sentiment Mix by Top Reviewed Products (%)', **TITLE_KW)
ax.set_xlabel(''); ax.set_ylabel('')
plt.yticks(rotation=0, fontsize=8)
plt.tight_layout()
plt.savefig('chart_5_sentiment_by_product.png', dpi=150, bbox_inches='tight')
plt.close()

print("\n✅  5 charts saved as individual PNG files.")

# ─────────────────────────────────────────────────────────────
# 7. CONCLUSIONS — PUBLIC OPINION & ACTIONABLE INSIGHTS
# ─────────────────────────────────────────────────────────────
pos_pct = sent_pct.get('Positive', 0)
neg_pct = sent_pct.get('Negative', 0)
neu_pct = sent_pct.get('Neutral', 0)
top_emotion = emotion_counts.most_common(1)[0][0]
named_products = product_sentiment_top.drop(index='Unknown product', errors='ignore')
worst_product = named_products['Negative'].idxmax() if 'Negative' in named_products and not named_products.empty else None
best_product = named_products['Positive'].idxmax() if 'Positive' in named_products and not named_products.empty else None

print("\n" + "=" * 60)
print("CONCLUSIONS")
print("=" * 60)
print(f"1. Public opinion is overwhelmingly positive: {pos_pct}% of reviews "
      f"express positive sentiment, vs. {neg_pct}% negative and {neu_pct}% neutral.")
print(f"2. '{top_emotion.capitalize()}' is the most frequently detected specific "
      f"emotion, appearing in {emotion_counts[top_emotion]:,} reviews.")
print(f"3. Text sentiment agrees with the star rating given {agreement:.1f}% of the "
      f"time, validating the lexicon-based approach as a reliable proxy for opinion.")
if best_product is not None:
    print(f"4. '{best_product[:50]}' has the strongest positive sentiment among "
          f"top-reviewed products ({product_sentiment_top.loc[best_product, 'Positive']:.0f}% positive).")
if worst_product is not None:
    print(f"5. '{worst_product[:50]}' has the highest share of negative sentiment "
          f"({product_sentiment_top.loc[worst_product, 'Negative']:.0f}% negative) "
          f"and is the clearest candidate for product or support investigation.")

# Save a JSON summary for the report-building step
summary = {
    "total_reviews": int(len(df)),
    "positive_pct": float(pos_pct), "neutral_pct": float(neu_pct), "negative_pct": float(neg_pct),
    "positive_count": int(sent_counts.get('Positive', 0)),
    "neutral_count": int(sent_counts.get('Neutral', 0)),
    "negative_count": int(sent_counts.get('Negative', 0)),
    "agreement_with_rating": float(round(agreement, 1)),
    "top_emotion": top_emotion,
    "top_emotion_count": int(emotion_counts[top_emotion]),
    "emotion_counts": {k: int(v) for k, v in emotion_counts.items()},
    "best_product": best_product, "worst_product": worst_product,
    "avg_word_count": float(round(df['word_count'].mean(), 1)),
}
with open('sentiment_summary.json', 'w') as f:
    json.dump(summary, f, indent=2)

# ─────────────────────────────────────────────────────────────
# 8. BUILD THE POLISHED WORD REPORT (.docx)
# ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARKGREEN = RGBColor(0x1B, 0x43, 0x32)
GREEN_RGB = RGBColor(0x2D, 0x6A, 0x4F)
GREY_RGB  = RGBColor(0x59, 0x59, 0x59)
HEADER_FILL, ROW_FILL = "2D6A4F", "EAF6EE"

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.bold = bold; run.font.name = 'Arial'
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = DARKGREEN if level == 1 else GREEN_RGB
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align_center=False):
    p = doc.add_paragraph()
    if align_center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.bold, run.italic = bold, italic
    run.font.size = Pt(size); run.font.name = 'Arial'
    if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text); run.bold = bold; run.font.name = 'Arial'
    return p

report = Document()
report.styles['Normal'].font.name = 'Arial'
report.styles['Normal'].font.size = Pt(11)

# Title
add_para(report, "Sentiment Analysis Report", bold=True, size=26, color=DARKGREEN, align_center=True)
add_para(report, "Amazon Product Reviews — Classifying Public Opinion", italic=True,
          size=14, color=GREEN_RGB, align_center=True)
add_para(report, "CodeAlpha Internship — Task 4  |  Author: Nesma Yahia", size=10,
          color=GREY_RGB, align_center=True)

# 1. How text was processed
add_heading(report, "1. How the Text Was Processed", level=1)
add_para(report,
    f"The dataset consists of {len(df):,} real Amazon product reviews (electronics "
    f"category — tablets, e-readers, and accessories), each including the review text, "
    f"a 1–5 star rating, and the product name. Rows missing review text or a star "
    f"rating were removed, since both are required to analyze and validate sentiment. "
    f"Each review was lightly cleaned — stripping URLs, HTML artifacts, and extra "
    f"whitespace — while deliberately preserving punctuation and capitalization "
    f"(e.g. \"AMAZING!!!\" vs. \"amazing\"), since the sentiment lexicon used in this "
    f"analysis treats those as meaningful signals of intensity, not noise to remove.")
add_para(report,
    f"On average, each review contains {summary['avg_word_count']} words — short "
    f"enough that a lexicon-based approach (rather than a large trained model) is both "
    f"appropriate and fast to run on the full dataset.")

# 2. How sentiment was assigned
add_heading(report, "2. How Sentiment Was Assigned", level=1)
add_para(report,
    "Sentiment was assigned using VADER (Valence Aware Dictionary and sEntiment "
    "Reasoner), a lexicon- and rule-based tool purpose-built for short, informal text "
    "like reviews and social media posts. VADER scores each piece of text from -1 "
    "(most negative) to +1 (most positive) using a 'compound' score, accounting for "
    "negation ('not good'), degree modifiers ('very good' vs. 'good'), punctuation "
    "emphasis, and capitalization.")
add_bullet(report, "compound score ≥ 0.05  →  classified as Positive")
add_bullet(report, "compound score ≤ -0.05  →  classified as Negative")
add_bullet(report, "compound score between -0.05 and 0.05  →  classified as Neutral")
add_para(report,
    f"To detect specific emotions beyond simple polarity, a compact keyword lexicon "
    f"(in the style of the NRC Emotion Lexicon) was applied, matching review text "
    f"against word lists for six emotions: joy, trust, anger, sadness, fear, and "
    f"surprise. A review can trigger more than one emotion.")
add_para(report,
    f"To validate the method, text-based sentiment was compared against the actual "
    f"star rating each reviewer gave (1–2 stars = negative, 3 = neutral, 4–5 = "
    f"positive). The two agreed on {summary['agreement_with_rating']}% of reviews, "
    f"giving confidence that the lexicon-based labels reflect genuine reader opinion "
    f"rather than noise.")

# 3. Sentiment counts table + chart
report.add_page_break()
add_heading(report, "3. Sentiment Classification Results", level=1)
table = report.add_table(rows=0, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
header = table.add_row()
for i, h in enumerate(['Sentiment', 'Review Count', 'Share of Total']):
    set_cell_text(header.cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), align_center=True)
    shade_cell(header.cells[i], HEADER_FILL)
for i, label in enumerate(['Positive', 'Neutral', 'Negative']):
    row = table.add_row()
    set_cell_text(row.cells[0], label, bold=True)
    set_cell_text(row.cells[1], f"{summary[label.lower()+'_count']:,}", align_center=True)
    set_cell_text(row.cells[2], f"{summary[label.lower()+'_pct']}%", align_center=True)
    if i % 2 == 0:
        for c in row.cells: shade_cell(c, ROW_FILL)

add_para(report, "")
img_p = report.add_paragraph(); img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.add_run().add_picture('chart_1_sentiment_counts.png', width=Inches(5.3))
img_p2 = report.add_paragraph(); img_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p2.add_run().add_picture('chart_2_sentiment_share.png', width=Inches(4.3))

# 4. Emotion detection
report.add_page_break()
add_heading(report, "4. Specific Emotions Detected", level=1)
add_para(report,
    f"Beyond positive/negative/neutral, the emotion lexicon surfaced more specific "
    f"signals in the review text. '{summary['top_emotion'].capitalize()}' was the "
    f"most frequently detected emotion, appearing in "
    f"{summary['top_emotion_count']:,} reviews — consistent with an overwhelmingly "
    f"positive review base, where customers express enthusiasm and confidence in "
    f"their purchase rather than just rating it favorably.")
img_p3 = report.add_paragraph(); img_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p3.add_run().add_picture('chart_3_emotion_counts.png', width=Inches(5.3))

# 5. Validation + product breakdown
report.add_page_break()
add_heading(report, "5. Validating Against Star Ratings", level=1)
add_para(report,
    f"Comparing text sentiment to the star rating each reviewer actually gave shows "
    f"strong alignment ({summary['agreement_with_rating']}% agreement) — high-star "
    f"reviews are reliably classified as positive, and low-star reviews as negative. "
    f"This cross-check matters: it confirms the sentiment labels are measuring real "
    f"opinion, not an artifact of the lexicon.")
img_p4 = report.add_paragraph(); img_p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p4.add_run().add_picture('chart_4_sentiment_vs_rating.png', width=Inches(5.3))

add_heading(report, "Sentiment by Product", level=2)
add_para(report,
    "Breaking sentiment down by product reveals which items are winning over "
    "customers and which ones are drawing complaints — actionable at the SKU level "
    "rather than just in aggregate.")
img_p5 = report.add_paragraph(); img_p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p5.add_run().add_picture('chart_5_sentiment_by_product.png', width=Inches(5.6))

# 6. Conclusions
report.add_page_break()
add_heading(report, "6. Conclusions: What This Means for the Business", level=1)
add_bullet(report,
    f"Public opinion is strongly positive: {summary['positive_pct']}% of reviews "
    f"are classified positive, vs. only {summary['negative_pct']}% negative — this "
    f"product line is well-received overall.", bold=True)
add_bullet(report,
    f"'{summary['top_emotion'].capitalize()}' dominates the emotional tone of "
    f"reviews, suggesting marketing messaging built around enthusiasm and "
    f"confidence (rather than just 'good ratings') will resonate with how "
    f"customers already talk about the product.", bold=True)
if summary['worst_product']:
    add_bullet(report,
        f"'{summary['worst_product'][:60]}' shows the highest share of negative "
        f"sentiment among top-reviewed products and is the clearest candidate for "
        f"a product-quality or customer-support review.", bold=True)
if summary['best_product']:
    add_bullet(report,
        f"'{summary['best_product'][:60]}' has the strongest positive sentiment "
        f"and could serve as a reference case — what is this product doing right "
        f"that others aren't?", bold=True)
add_bullet(report,
    f"Because text sentiment and star ratings agree {summary['agreement_with_rating']}% "
    f"of the time, this lexicon-based pipeline is reliable enough to monitor "
    f"ongoing review streams automatically, flagging sudden spikes in negative "
    f"sentiment or specific emotions like anger or sadness before they show up "
    f"as a ratings drop.", bold=True)
add_para(report,
    "Together, these results turn thousands of individual reviews into a small set "
    "of clear signals: what customers feel, how strongly, and which products need "
    "attention — the kind of pattern that's invisible in a spreadsheet of star "
    "ratings alone but obvious once the text itself is analyzed.")

report_path = 'Amazon_Sentiment_Analysis_Report.docx'
report.save(report_path)

# Patch the zoom attribute so the docx passes strict OOXML validation
import zipfile, shutil, tempfile
def fix_zoom_attribute(path):
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = data.replace(b'<w:zoom w:val="bestFit"/>',
                                     b'<w:zoom w:val="bestFit" w:percent="100"/>')
            zout.writestr(item, data)
    shutil.move(tmp_path, path)
fix_zoom_attribute(report_path)

print(f"\n✅  Word report saved to '{report_path}'")
