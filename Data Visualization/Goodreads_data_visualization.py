"""
=============================================================
  Data Visualization Task — Goodreads Books Dataset
  Goal: Turn raw numbers into charts that are instantly
        understandable, and tell the story behind each one.
=============================================================

This script:
  1. Loads and lightly cleans the Goodreads books dataset
  2. Builds five chart types, each chosen for what it's best at:
       - Bar chart   -> comparing categories (top authors)
       - Line chart  -> trend over time (books published per year)
       - Pie chart   -> proportions, used sparingly (language mix)
       - Histogram   -> distribution of values (average rating)
       - Heatmap     -> correlation between variables
  3. Saves each chart as its own PNG (for a portfolio) plus prints
     a short data story for each one, explaining what it means.
"""

import pandas as pd
import matplotlib
matplotlib.use('Agg')                 # headless rendering
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
import numpy as np
import warnings
warnings.filterwarnings('ignore')

# ─────────────────────────────────────────────────────────────
# 0. STYLE — consistent, portfolio-friendly look
# ─────────────────────────────────────────────────────────────
sns.set_style("whitegrid")
plt.rcParams.update({
    'font.family': 'DejaVu Sans',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'figure.facecolor': 'white',
    'axes.facecolor': 'white',
})
PRIMARY   = '#2D6A4F'
ACCENT    = '#FF6B6B'
PALETTE   = sns.color_palette('YlGn', 10)[::-1]
TITLE_KW  = dict(fontsize=14, fontweight='bold', color='#1B4332', pad=14)
LABEL_KW  = dict(fontsize=11, color='#333333')

# ─────────────────────────────────────────────────────────────
# 1. LOAD & CLEAN DATA
# ─────────────────────────────────────────────────────────────
df = pd.read_csv('books.csv', on_bad_lines='skip')
df.columns = df.columns.str.strip()

# Handle missing data: fill non-essential text fields, drop rows
# missing the numbers we actually need to chart
for col in ['publisher', 'language_code', 'authors', 'title']:
    if col in df.columns:
        df[col] = df[col].fillna('Unknown')

essential = ['average_rating', 'num_pages', 'ratings_count',
             'text_reviews_count', 'publication_date']
df = df.dropna(subset=[c for c in essential if c in df.columns])

# Remove invalid placeholder rows and parse dates
df = df[(df['num_pages'] > 0) & (df['average_rating'] > 0)]
df['pub_year'] = pd.to_datetime(df['publication_date'], errors='coerce').dt.year

print(f"Loaded and cleaned dataset: {len(df):,} books ready for visualization.\n")

# ─────────────────────────────────────────────────────────────
# 2. BAR CHART — Comparing categories: Top 10 Authors
# ─────────────────────────────────────────────────────────────
top_authors = df['authors'].value_counts().head(10)

fig, ax = plt.subplots(figsize=(9, 6))
bars = ax.barh(top_authors.index[::-1], top_authors.values[::-1],
                color=PALETTE, edgecolor='#1B4332', linewidth=0.6)
for bar, val in zip(bars, top_authors.values[::-1]):
    ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
            str(val), va='center', fontsize=10, color='#1B4332')
ax.set_title('Who Are the Most Prolific Authors?', **TITLE_KW)
ax.set_xlabel('Number of Books', **LABEL_KW)
plt.tight_layout()
plt.savefig('chart_1_bar_top_authors.png', dpi=150, bbox_inches='tight')
plt.close()

tie_note = " (tied)" if top_authors.values[0] == top_authors.values[1] else ""
story_bar = (
    f"{top_authors.index[0]} leads the dataset with {top_authors.values[0]} books"
    f"{', tied with ' + top_authors.index[1] if tie_note else ', followed by ' + top_authors.index[1] + f' ({top_authors.values[1]})'}. "
    f"A bar chart makes this comparison instant: instead of scanning a table of "
    f"names and counts, the reader sees at a glance who dominates the catalogue."
)
print("CHART 1 — Bar Chart (Top Authors)")
print(story_bar, "\n")

# ─────────────────────────────────────────────────────────────
# 3. LINE CHART — Trend over time: Books Published Per Year
# ─────────────────────────────────────────────────────────────
yearly = df[df['pub_year'].between(1980, 2010)]['pub_year'].value_counts().sort_index()
peak_year = int(yearly.idxmax())

fig, ax = plt.subplots(figsize=(10, 5.5))
ax.plot(yearly.index, yearly.values, color=PRIMARY, lw=2.5, marker='o', markersize=4)
ax.fill_between(yearly.index, yearly.values, alpha=0.15, color=PRIMARY)
ax.annotate(f'Peak: {peak_year} ({yearly.max()} books)',
            xy=(peak_year, yearly.max()),
            xytext=(peak_year - 12, yearly.max() * 0.8),
            arrowprops=dict(arrowstyle='->', color=ACCENT, lw=1.5),
            color=ACCENT, fontsize=10, fontweight='bold')
ax.set_title('Publishing Activity Over Time (1980–2010)', **TITLE_KW)
ax.set_xlabel('Year', **LABEL_KW); ax.set_ylabel('Number of Books Published', **LABEL_KW)
plt.tight_layout()
plt.savefig('chart_2_line_publications_over_time.png', dpi=150, bbox_inches='tight')
plt.close()

story_line = (
    f"Publishing volume in this dataset climbs steadily from the 1980s and peaks in "
    f"{peak_year} at {yearly.max()} books, before falling off sharply. A line chart is "
    f"the right tool here because the story is about change over time — the eye "
    f"follows the slope and immediately sees growth, peak, and decline, which a table "
    f"of yearly counts would hide. (The drop after the peak likely reflects when this "
    f"dataset snapshot was collected, not an actual publishing slowdown.)"
)
print("CHART 2 — Line Chart (Publications Over Time)")
print(story_line, "\n")

# ─────────────────────────────────────────────────────────────
# 4. PIE CHART — Proportions, used sparingly: Language Mix
# ─────────────────────────────────────────────────────────────
lang = df['language_code'].value_counts()
top_n = 5
pie_values = list(lang.head(top_n).values) + [lang.iloc[top_n:].sum()]
pie_labels = list(lang.head(top_n).index) + ['Other']
dom_pct = pie_values[0] / pie_values[:].__len__() if False else (lang.iloc[0] / len(df)) * 100

fig, ax = plt.subplots(figsize=(8, 7))
colors = sns.color_palette('YlGn', len(pie_values))
wedges, _, autotexts = ax.pie(
    pie_values, autopct=lambda p: f'{p:.1f}%' if p >= 4 else '', startangle=140,
    colors=colors, explode=[0.05] + [0.02] * (len(pie_values) - 1),
    textprops={'fontsize': 12, 'color': '#1B4332', 'fontweight': 'bold'},
    wedgeprops={'edgecolor': 'white', 'linewidth': 1.5})
ax.legend(wedges, [f'{lbl}  ({val/sum(pie_values)*100:.1f}%)' for lbl, val in zip(pie_labels, pie_values)],
          title='Language', loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=10, frameon=False)
ax.set_title('What Languages Are Books Published In?', **TITLE_KW)
plt.tight_layout()
plt.savefig('chart_3_pie_language_share.png', dpi=150, bbox_inches='tight')
plt.close()

story_pie = (
    f"English ('{lang.index[0]}') accounts for {dom_pct:.1f}% of all books, making it "
    f"by far the dominant language in the catalogue. A pie chart works well here "
    f"because the question is simple — 'what share does each language hold?' — and "
    f"there are few enough categories (6) that the slices stay readable. Pie charts "
    f"are used sparingly in this project precisely because they get cluttered fast "
    f"with more categories than that."
)
print("CHART 3 — Pie Chart (Language Share)")
print(story_pie, "\n")

# ─────────────────────────────────────────────────────────────
# 5. HISTOGRAM — Distribution of values: Average Rating
# ─────────────────────────────────────────────────────────────
mean_r, median_r = df['average_rating'].mean(), df['average_rating'].median()

fig, ax = plt.subplots(figsize=(9, 5.5))
n, bins, patches = ax.hist(df['average_rating'], bins=40, edgecolor='white', linewidth=0.5)
for patch, val in zip(patches, bins):
    patch.set_facecolor(plt.cm.YlGn(0.25 + 0.65 * (val - 2) / 3))
ax.axvline(mean_r, color=ACCENT, ls='--', lw=2, label=f'Mean: {mean_r:.2f}')
ax.axvline(median_r, color='#FFB100', ls='--', lw=2, label=f'Median: {median_r:.2f}')
ax.legend(fontsize=10)
ax.set_title('How Are Book Ratings Distributed?', **TITLE_KW)
ax.set_xlabel('Average Rating (0–5)', **LABEL_KW); ax.set_ylabel('Number of Books', **LABEL_KW)
plt.tight_layout()
plt.savefig('chart_4_histogram_rating_distribution.png', dpi=150, bbox_inches='tight')
plt.close()

story_hist = (
    f"Most books cluster tightly between roughly 3.7 and 4.2 stars, with a mean of "
    f"{mean_r:.2f} and a median of {median_r:.2f} that sit almost on top of each "
    f"other — a sign of a fairly symmetric, slightly left-skewed distribution rather "
    f"than a few outliers pulling the average around. A histogram is the right chart "
    f"for this because the question isn't about individual books, it's about the "
    f"overall shape of ratings across the whole dataset."
)
print("CHART 4 — Histogram (Rating Distribution)")
print(story_hist, "\n")

# ─────────────────────────────────────────────────────────────
# 6. HEATMAP — Correlation between variables
# ─────────────────────────────────────────────────────────────
corr_cols = ['average_rating', 'num_pages', 'ratings_count', 'text_reviews_count']
corr_labels = ['Avg Rating', 'Pages', 'Ratings Count', 'Text Reviews']
corr = df[corr_cols].corr()

fig, ax = plt.subplots(figsize=(7, 6))
sns.heatmap(corr, annot=True, fmt='.2f', cmap='YlGn', vmin=-0.1, vmax=1,
            xticklabels=corr_labels, yticklabels=corr_labels,
            linewidths=1, linecolor='white', cbar_kws={'label': 'Correlation'}, ax=ax)
ax.set_title('Which Variables Move Together?', **TITLE_KW)
plt.tight_layout()
plt.savefig('chart_5_heatmap_correlations.png', dpi=150, bbox_inches='tight')
plt.close()

ratings_reviews_corr = corr.loc['ratings_count', 'text_reviews_count']
rating_pages_corr = corr.loc['average_rating', 'num_pages']
story_heat = (
    f"The heatmap reveals one standout relationship: ratings count and text reviews "
    f"count are strongly correlated ({ratings_reviews_corr:.2f}) — books that get "
    f"rated a lot also get reviewed a lot, which makes sense, since both come from "
    f"engaged readers. Everything else is weak, including average rating versus page "
    f"count ({rating_pages_corr:.2f}), meaning a book's length tells you almost "
    f"nothing about how well it's rated. A heatmap is ideal here because it lets the "
    f"reader scan every pairwise relationship at once instead of reading four "
    f"separate correlation numbers out of context."
)
print("CHART 5 — Heatmap (Correlations)")
print(story_heat, "\n")

# ─────────────────────────────────────────────────────────────
# 7. SAVE THE DATA STORY (used later to build the PDF/DOCX report)
# ─────────────────────────────────────────────────────────────
import json
data_story = {
    "summary": {
        "total_books": int(len(df)),
        "mean_rating": round(float(mean_r), 2),
        "median_rating": round(float(median_r), 2),
        "top_author": top_authors.index[0],
        "top_author_count": int(top_authors.values[0]),
        "peak_year": peak_year,
        "peak_year_books": int(yearly.max()),
        "dominant_language": lang.index[0],
        "dominant_language_pct": round(float(dom_pct), 1),
        "ratings_reviews_corr": round(float(ratings_reviews_corr), 3),
        "rating_pages_corr": round(float(rating_pages_corr), 3),
    },
    "stories": {
        "bar": story_bar,
        "line": story_line,
        "pie": story_pie,
        "histogram": story_hist,
        "heatmap": story_heat,
    }
}
with open('data_story.json', 'w') as f:
    json.dump(data_story, f, indent=2)

print("=" * 60)
print("All 5 charts saved as individual PNG files.")
print("Narrative summary saved to 'data_story.json'.")
print("=" * 60)

# ─────────────────────────────────────────────────────────────
# 8. BUILD THE POLISHED WORD REPORT (.docx)
# ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARKGREEN = RGBColor(0x1B, 0x43, 0x32)
GREEN_RGB = RGBColor(0x2D, 0x6A, 0x4F)
GREY_RGB  = RGBColor(0x59, 0x59, 0x59)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = DARKGREEN if level == 1 else GREEN_RGB
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align_center=False):
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = color
    return p

def add_chart_section(doc, number, title, chart_kind, image_path, narrative):
    add_heading(doc, f"{number}. {title}", level=1)
    add_para(doc, f"Chart type: {chart_kind}", italic=True, size=10, color=GREY_RGB)
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(image_path, width=Inches(5.6))
    add_heading(doc, "What this tells us", level=2)
    add_para(doc, narrative)

report = Document()
report.styles['Normal'].font.name = 'Arial'
report.styles['Normal'].font.size = Pt(11)

# Title
add_para(report, "Data Visualization Report", bold=True, size=26, color=DARKGREEN, align_center=True)
add_para(report, "Goodreads Books Dataset — Turning Numbers Into Charts", italic=True,
          size=14, color=GREEN_RGB, align_center=True)
add_para(report, "From raw data to a clear, decision-supporting visual story.",
          size=10, color=GREY_RGB, align_center=True)

# Introduction
add_heading(report, "Introduction", level=1)
add_para(report,
    f"This report transforms {len(df):,} rows of raw Goodreads book data into five "
    f"purpose-built visuals — a bar chart, a line chart, a pie chart, a histogram, and a "
    f"heatmap — each chosen because it is the clearest way to answer a specific question "
    f"about the data. Rather than just presenting charts, each section explains what the "
    f"chart shows and why it matters, in the same way a chart would support a real "
    f"business or editorial decision.")

# Five chart sections, each on (mostly) its own page
add_chart_section(report, 1, "Who Are the Most Prolific Authors?",
                   "Bar Chart — best for comparing categories",
                   "chart_1_bar_top_authors.png", story_bar)

report.add_page_break()
add_chart_section(report, 2, "Publishing Activity Over Time",
                   "Line Chart — best for trends over time",
                   "chart_2_line_publications_over_time.png", story_line)

report.add_page_break()
add_chart_section(report, 3, "What Languages Are Books Published In?",
                   "Pie Chart — best for simple proportions, used sparingly",
                   "chart_3_pie_language_share.png", story_pie)

report.add_page_break()
add_chart_section(report, 4, "How Are Book Ratings Distributed?",
                   "Histogram — best for the distribution/shape of a value",
                   "chart_4_histogram_rating_distribution.png", story_hist)

report.add_page_break()
add_chart_section(report, 5, "Which Variables Move Together?",
                   "Heatmap — best for correlations between variables",
                   "chart_5_heatmap_correlations.png", story_heat)

# Data story / conclusion
report.add_page_break()
add_heading(report, "The Story in Plain English", level=1)
add_para(report,
    f"Put together, these five charts tell a single coherent story about the Goodreads "
    f"catalogue. The collection is overwhelmingly English-language ({dom_pct:.1f}%), "
    f"concentrated among a small set of highly prolific authors led by "
    f"{top_authors.index[0]} ({top_authors.values[0]} books), and skewed toward titles "
    f"published in the years leading up to {peak_year}, when the dataset's coverage "
    f"peaks at {yearly.max():,} books in a single year.")
add_para(report,
    f"On quality and engagement: books are rated consistently well, clustering around "
    f"{mean_r:.2f}/5, and that rating has almost no relationship to how long a book is "
    f"(correlation = {rating_pages_corr:.3f}). What does matter is reader engagement — "
    f"books that accumulate many ratings also accumulate many written reviews "
    f"(correlation = {ratings_reviews_corr:.3f}), confirming that ratings count is a "
    f"reliable proxy for overall reader engagement, not just popularity in isolation.")
add_heading(report, "Why this matters for decisions", level=2)
add_para(report,
    "For a publisher or platform, this points to a clear takeaway: marketing and "
    "discovery efforts built around 'ratings count' will naturally also surface highly "
    "reviewed books, since the two move together. Meanwhile, page count is not a useful "
    "signal for predicting reader satisfaction and shouldn't be used to filter or "
    "recommend titles. Visualizing this — rather than only stating the correlation "
    "numbers — makes the pattern immediately convincing and easy to act on.")

report_path = 'Goodreads_Data_Visualization_Report.docx'
report.save(report_path)

# python-docx's default template can omit the required w:percent attribute on
# <w:zoom>, which some strict OOXML validators reject. Patch it in-place.
import zipfile, shutil, tempfile

def fix_zoom_attribute(path):
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = data.replace(b'<w:zoom w:val="bestFit"/>',
                                     b'<w:zoom w:val="bestFit" w:percent="100"/>')
            zout.writestr(item, data)
    shutil.move(tmp_path, path)

fix_zoom_attribute(report_path)
print(f"✅  Word report saved to '{report_path}'")
